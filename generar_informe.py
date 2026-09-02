# -*- coding: utf-8 -*-
"""Genera el informe del Proyecto 1 como .docx.

Formato según CONVENCION_GENERAL_DOCX.md y redaccion según Mi_Lenguaje.md.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

Raíz = Path(r"d:/Tareas/Computacion Paralela/informe-proyecto-1-paralela")
ANEXOS = Raíz / "anexos"
FUENTE = "Times New Roman"

doc = Document()

# --- Margenes de una pulgada -------------------------------------------------
for s in doc.sections:
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

def fijar_fuente(estilo, tam, negrilla=False, color=None):
    estilo.font.name = FUENTE
    estilo.font.size = Pt(tam)
    estilo.font.bold = negrilla
    if color is not None:
        estilo.font.color.rgb = color
    rpr = estilo.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), FUENTE)

est = doc.styles
fijar_fuente(est['Normal'], 11)
est['Normal'].paragraph_format.line_spacing = 1.15
est['Normal'].paragraph_format.space_after = Pt(6)
fijar_fuente(est['Heading 1'], 14, True, RGBColor(0, 0, 0))
est['Heading 1'].paragraph_format.space_before = Pt(16)
est['Heading 1'].paragraph_format.space_after = Pt(8)
fijar_fuente(est['Heading 2'], 12, True, RGBColor(0, 0, 0))
est['Heading 2'].paragraph_format.space_before = Pt(12)
est['Heading 2'].paragraph_format.space_after = Pt(6)
for n in ('List Bullet', 'List Number', 'Caption'):
    try:
        fijar_fuente(est[n], 11)
    except KeyError:
        pass

# --- Utilidades --------------------------------------------------------------
def p(texto="", tam=11, negrilla=False, alineacion=None, cursiva=False,
      espacio_antes=None, sangria=None):
    par = doc.add_paragraph()
    if alineacion is not None:
        par.alignment = alineacion
    if espacio_antes is not None:
        par.paragraph_format.space_before = Pt(espacio_antes)
    if sangria is not None:
        par.paragraph_format.left_indent = Inches(sangria)
        par.paragraph_format.right_indent = Inches(sangria)
    if texto:
        r = par.add_run(texto)
        r.font.name = FUENTE; r.font.size = Pt(tam)
        r.bold = negrilla; r.italic = cursiva
    return par

def _forzar_fuente(par, tam):
    # El tema del documento pisa la fuente del estilo, entonces se fija en el run.
    for r in par.runs:
        r.font.name = FUENTE
        r.font.size = Pt(tam)
        r.bold = True
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rf.set(qn(a), FUENTE)
    return par

def h1(t):
    return _forzar_fuente(doc.add_heading(t, level=1), 14)
def h2(t):
    return _forzar_fuente(doc.add_heading(t, level=2), 12)

def vineta(texto, tam=11):
    par = doc.add_paragraph(style='List Bullet')
    r = par.add_run(texto)
    r.font.name = FUENTE; r.font.size = Pt(tam)
    par.paragraph_format.space_after = Pt(3)
    return par

CONT = {"tabla": 0, "figura": 0}

def pie(clase, texto):
    CONT[clase] += 1
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after = Pt(12)
    etiqueta = ("Tabla " if clase == "tabla" else "Figura ") + str(CONT[clase]) + ". "
    r1 = par.add_run(etiqueta); r1.bold = True
    r1.font.name = FUENTE; r1.font.size = Pt(10)
    r2 = par.add_run(texto); r2.font.name = FUENTE; r2.font.size = Pt(10)
    return par

def figura(archivo, texto, ancho=6.0):
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(ANEXOS / archivo), width=Inches(ancho))
    pie("figura", texto)

def tabla(cabeceras, filas, texto, anchos=None):
    t = doc.add_table(rows=1, cols=len(cabeceras))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, c in enumerate(cabeceras):
        celda = t.rows[0].cells[i]
        celda.text = ""
        par = celda.paragraphs[0]; par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = par.add_run(c); r.bold = True
        r.font.name = FUENTE; r.font.size = Pt(10)
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = ""
            par = celdas[i].paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = par.add_run(str(v)); r.font.name = FUENTE; r.font.size = Pt(10)
    if anchos:
        for fila in t.rows:
            for i, an in enumerate(anchos):
                fila.cells[i].width = Inches(an)
    pie("tabla", texto)
    return t

def cita_textual(texto, atribucion):
    par = p(texto, tam=11, cursiva=True, sangria=0.5)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par2 = p(atribucion, tam=10, sangria=0.5)
    par2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par2.paragraph_format.space_after = Pt(12)

def salto():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================================
#  PORTADA
# ============================================================================
for _ in range(3):
    p()
p("Universidad del Valle de Guatemala", 12, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
p("Facultad de Ingeniería", 12, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
p("Computación Paralela y Distribuida", 12, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
p("Semestre 2, 2026", 12, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    p()
p("Paralelización con OpenMP de un protector de pantalla", 18, True,
  WD_ALIGN_PARAGRAPH.CENTER)
p("de terreno por voxeles", 18, True, WD_ALIGN_PARAGRAPH.CENTER)
p()
p("Proyecto 1", 14, True, WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(4):
    p()
p("José Galindo", 12, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
p("Joel Jaquez", 12, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    p()
p("Ing. Santiago Solórzano", 11, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
p()
p("Guatemala, 2 de septiembre de 2026", 11, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
p()
p("Repositorio del código fuente", 11, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
p("https://github.com/DerNait/minecraft-screen-saver", 11,
  alineacion=WD_ALIGN_PARAGRAPH.CENTER)
salto()

# ============================================================================
#  Índice
# ============================================================================
h1("Índice")
par = doc.add_paragraph()
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), r'TOC \o "1-2" \h \z \u')
sub = OxmlElement('w:r'); subt = OxmlElement('w:t')
subt.text = "Para ver el índice, seleccione esta línea y presione F9 en Word."
sub.append(subt); fld.append(sub)
par._p.append(fld)
salto()

# ============================================================================
#  Introducción
# ============================================================================
h1("Introducción")
p("Este informe documenta el diseño, la implementación y la medición de un protector "
  "de pantalla que dibuja un terreno de bloques al estilo de Minecraft, primero en una "
  "versión secuencial y después en una versión paralela hecha con OpenMP. OpenMP es un "
  "conjunto de directivas que se agregan al código en C o C++ para repartir trabajo "
  "entre varios hilos que comparten la misma memoria.")
p("El programa recibe un parámetro N que indica cuántos bloques debe generar y "
  "renderizar. A partir de una semilla arma un mundo completo, lo construye capa por "
  "capa de abajo hacia arriba, lo sostiene unos segundos, lo desarma de arriba hacia "
  "abajo y vuelve a empezar con un mundo nuevo. Cada bloque tiene posición, estado y "
  "física propia, entonces el trabajo por fotograma crece de forma lineal con N.")
p("El objetivo del trabajo fue tomar esa versión secuencial y acelerarla repartiendo "
  "el trabajo entre hilos, midiendo después el speedup y la eficiencia que se lograron. "
  "El speedup es cuántas veces más rápido corre la versión paralela comparada con la "
  "secuencial, y la eficiencia es ese speedup dividido entre la cantidad de hilos que "
  "se usaron.")
p("El resultado principal fue que con dos millones de bloques la versión secuencial "
  "promedia 53 fotogramas por segundo y deja 264 de 300 fotogramas medidos por debajo "
  "del límite de 60, mientras que la versión paralela sostiene 117 fotogramas por "
  "segundo y prácticamente ninguno queda por debajo de ese límite.")
salto()

# ============================================================================
#  ANTECEDENTES
# ============================================================================
h1("Antecedentes")
p("La programación paralela con memoria compartida parte de una idea sencilla. Si un "
  "problema se puede partir en pedazos que no dependen entre si, varios hilos pueden "
  "trabajar cada uno en su pedazo al mismo tiempo. OpenMP facilita ese reparto porque "
  "permite ir modificando un programa secuencial poco a poco, agregando directivas "
  "sobre los bucles que conviene repartir, sin tener que reescribir todo el programa "
  "desde cero (Chapman, Jost y van der Pas, 2007).")
p("Para decidir qué se reparte y cómo, el curso trabaja el método PCAM propuesto por "
  "Foster (1995). El método tiene cuatro pasos. Particionar consiste en partir el "
  "problema en las unidades de trabajo más pequenas posibles. Comunicar consiste en "
  "identificar qué dependencias existen entre esas unidades. Aglomerar consiste en "
  "juntar unidades pequenas en grupos más grandes para que el costo de repartir no se "
  "coma la ganancia. Mapear consiste en asignar esos grupos a los hilos disponibles.")
p("Ahora bien, la mejora que se puede conseguir tiene un techo. Amdahl (1967) mostró "
  "que la parte del programa que no se puede repartir termina limitando la ganancia "
  "total, por muchos procesadores que se agreguen. En su artículo original lo plantea "
  "así cuando habla del trabajo de administración de datos.")
cita_textual(
  "\"The nature of this overhead appears to be sequential so that it is unlikely to be "
  "amenable to parallel processing techniques.\"",
  "Amdahl, 1967, p. 483")
p("Esa observación aplica directo a este proyecto. Una parte del tiempo de cada "
  "fotograma se va en dibujar, y ese dibujo lo hace la tarjeta gráfica de la misma "
  "forma en las dos versiones. Esa parte se mantiene igual sin importar cuántos hilos "
  "se usen, entonces marca el límite de lo que se puede mejorar.")
p("En cuanto al problema elegido, un protector de pantalla de bloques encaja bien "
  "porque cumple dos condiciones. La primera es que la cantidad de elementos se "
  "controla con un parámetro, lo que permite probar cargas de distinto tamaño. La "
  "segunda es que dentro de un mismo fotograma cada bloque se actualiza usando "
  "solamente sus propios datos, entonces no hace falta que los hilos se coordinen para "
  "avanzar la simulación.")

# ============================================================================
#  CUERPO
# ============================================================================
h1("Diseño del protector de pantalla")
p("El programa genera un mundo cuadrado de bloques a partir de una semilla. Esa "
  "semilla decide el bioma, que puede ser llanura, bosque, abedules, desierto o taiga, "
  "y con el bioma cambian los colores, los tipos de bloque y la vegetación. El mundo "
  "pasa por cuatro fases que se repiten en ciclo. Primero se construye subiendo capa "
  "por capa, después se sostiene armado unos segundos, luego se desarma de arriba "
  "hacia abajo y al final hay una pausa antes de generar un mundo nuevo.")
figura("screensaver_mundo_armado.png",
       "El protector de pantalla con el mundo completamente armado en el bioma de bosque.",
       5.4)
p("El enunciado del proyecto pide seis condiciones y el programa las cumple todas. "
  "Recibe el parámetro N que indica la cantidad de bloques a renderizar, con un rango "
  "aceptado de 64 hasta 20 millones. Despliega varios colores generados de forma "
  "pseudoaleatoria a partir de la semilla. Usa un lienzo de 1280 por 720 píxeles, que "
  "está por encima del mínimo de 640 por 480 y se puede cambiar por argumento. Tiene "
  "movimiento, porque los bloques caen y la cámara recorre el mundo. Incorpora física "
  "y trigonometría, con gravedad en la caída de cada bloque y funciones de seno y "
  "coseno en la cámara y en la dirección de los animales. Y despliega los cuadros por "
  "segundo en el título de la ventana.")
p("Vale la pena notar que N funciona como un objetivo y no como una cantidad exacta. "
  "El programa deduce de N el lado de la retícula cuadrada que produce aproximadamente "
  "esa cantidad de bloques y luego genera el terreno completo. Con N igual a dos "
  "millones el mundo real termina con 1 918 719 bloques.")

h2("Los cálculos que se hacen por cada elemento")
p("Cada bloque recorre cinco estados durante el ciclo, que son Waiting, Falling, "
  "Placed, Vanishing y Gone. La función updateWorld recorre los N bloques en cada "
  "fotograma y avanza ese estado.")
p("El cálculo más importante ocurre mientras el bloque cae. Se aplica una integración "
  "de Euler de la gravedad, que consiste en actualizar primero la velocidad sumandole "
  "la aceleración por el paso de tiempo, y después actualizar la posición restandole "
  "esa velocidad nueva por el mismo paso de tiempo. La aceleración usada es de 45 "
  "bloques por segundo al cuadrado y cada bloque cae desde 6 bloques arriba de su "
  "posición final. Al mismo tiempo el bloque crece de escala 0 a 1 en 0.12 segundos, "
  "que es lo que da el efecto de aparición.")
p("La segunda función que corre sobre todos los elementos es buildInstanceBuffer, que "
  "decide cuáles bloques vale la pena mandar a la tarjeta gráfica. Un bloque se "
  "descarta si ya desaparecio, o si sus seis vecinos están ocupados, porque en ese "
  "caso ninguna de sus caras se puede ver desde afuera. El efecto de ese descarte es "
  "grande. Con dos millones de bloques generados, en pantalla se dibujan alrededor de "
  "309 mil, o sea cerca de una sexta parte.")
figura("descarte_invisibles.png",
       "Corte del terreno donde se ve que solo la cáscara exterior llega a dibujarse.",
       6.0)

h1("Aplicación del método PCAM")
p("Antes de escribir una sola directiva se aplicó el método PCAM sobre el programa "
  "secuencial ya funcionando.")
p("En el paso de particionar se identificó cuál es la unidad de trabajo más pequeña es una "
  "columna del plano horizontal del mundo, o un bloque de la lista, según la etapa. "
  "Una columna es una posición en el plano con todos los bloques que se apilan encima "
  "de ella. Un mundo de 380 por 380 tiene 144 400 columnas.")
p("En el paso de comunicar se buscaron las dependencias reales entre esas unidades y "
  "resultaron ser pocas. El mapa de alturas y los estratos no tienen ninguna, porque "
  "cada columna escribe únicamente en si misma. Aparecieron dos dependencias que si "
  "importan. La primera es que la copa de un árbol nace en una columna pero ocupa "
  "celdas de las columnas vecinas. La segunda es que el descarte de bloques invisibles "
  "necesita leer el estado ya actualizado de los seis vecinos de cada bloque.")
p("En el paso de aglomerar se agruparon las unidades para que el costo de repartir no "
  "domine. Se juntaron en rangos contiguos de iteraciones para el terreno, y en chunks "
  "rectangulares para la vegetación. Un chunk es un rectángulo de columnas asignado a "
  "un hilo. Con 12 hilos la cuadrícula queda de 4 por 3, y cada chunk agrupa alrededor "
  "de 12 mil columnas.")
p("En el paso de mapear se asignaron esos grupos a los hilos con reparto estático, o "
  "sea repartiendo todo antes de empezar en lugar de ir entregando trabajo sobre la "
  "marcha.")

h1("Implementación paralela con OpenMP")
p("El programa tiene siete regiones paralelas. Cuatro están en la generación del "
  "mundo, que ocurre una sola vez por mundo, y tres están dentro del bucle que se "
  "repite en cada fotograma. La Tabla 1 muestra cada una con la directiva que usa y la "
  "razón de esa elección.")
tabla(
  ["Región", "Directiva y cláusulas", "Por que así"],
  [
    ["Mapa de alturas",
     "parallel for collapse(2) schedule(static) reduction(max)",
     "collapse(2) fusiona los dos bucles anidados y pasa de 380 unidades de trabajo a 144 400. static porque el costo por columna es parejo. reduction obtiene la altura máxima sin carrera."],
    ["Estratos del terreno",
     "parallel for collapse(2) schedule(static)",
     "Cada columna escribe solo sus propias celdas, entonces no hay conflicto posible."],
    ["Vegetación",
     "parallel con reparto manual en chunks 2D",
     "Una copa de árbol cruza fronteras, entonces hace falta un dominio con forma y no un rango de indices."],
    ["Lista de bloques",
     "parallel con barrier, single y suma de prefijos",
     "La cantidad de bloques por hilo es variable. El prefijo asigna rangos exclusivos y conserva el orden de la versión secuencial."],
    ["Física de los bloques",
     "parallel for schedule(static) reduction(+)",
     "reduction cuenta los bloques vivos. La barrera implícita al cerrar el for ordena esta etapa antes de la siguiente."],
    ["Descarte y empaquetado",
     "parallel con barrier, single y suma de prefijos",
     "Otra compactación de tamaño variable. Además se guarda la marca de visibilidad para no repetir las seis lecturas por bloque."],
    ["Bloques que dan sombra",
     "parallel con barrier, single y suma de prefijos",
     "Mismo patrón. El single redimensiona el vector de salida antes de que los hilos escriban."],
  ],
  "Regiones paralelas del programa con la directiva que usa cada una.",
  anchos=[1.4, 2.0, 3.0])

h2("El halo de la vegetación")
p("El caso más difícil fue la vegetación. Un árbol tiene su tronco en una columna, "
  "pero su copa ocupa celdas hasta dos posiciones alrededor. Si el reparto se hiciera "
  "por rangos de indices, dos hilos podrían terminar escribiendo la misma celda al "
  "mismo tiempo, lo que produce una condición de carrera. Una condición de carrera "
  "ocurre cuando dos hilos tocan la misma posición de memoria, al menos uno escribe, y "
  "no hay un orden garantizado entre ellos.")
p("La solución fue darle a cada hilo dos regiones distintas. La región de lectura es "
  "su chunk más dos celdas hacia afuera en las cuatro direcciones, y a ese margen se "
  "le llama halo. La región de escritura es únicamente su propio chunk. Entonces el "
  "hilo revisa árboles cuya raíz pertenece al vecino, calcula la copa completa y "
  "escribe solamente los pedazos que caen dentro de su territorio. El hilo vecino hace "
  "lo mismo con ese árbol y escribe la otra mitad. Entre los dos reconstruyen el árbol "
  "sin haberse comunicado.")
figura("halo_vegetacion.png",
       "Región que lee y región que escribe cada hilo al colocar la vegetación.",
       6.0)
p("El halo mide dos celdas porque ese es el radio horizontal máximo que alcanza una "
  "copa en el código actual. El costo de esta solución es que los bordes entre chunks "
  "se calculan dos veces, una por cada hilo vecino. Ese trabajo repetido se paga a "
  "cambio de no necesitar ninguna sincronización, y como todo se deriva de la misma "
  "semilla los dos hilos calculan exactamente lo mismo.")

h1("Protección de la memoria compartida y sincronía")
p("En las siete regiones paralelas no hay ninguna sección critical, ningún atomic y "
  "ningún lock. La escritura compartida se eliminó por diseño en lugar de protegerse. "
  "Cada hilo escribe en un rango de memoria que ningún otro hilo toca.")
p("Cuando la cantidad de resultados que produce cada hilo es variable, como pasa en el "
  "descarte de bloques invisibles, se usa un patrón de compactación en tres fases. En "
  "la primera fase cada hilo recorre su rango y cuenta cuántos elementos pasan el "
  "filtro. En la segunda fase un solo hilo, dentro de una región single, acumula esos "
  "conteos y le asigna a cada hilo el índice donde debe empezar a escribir. En la "
  "tercera fase cada hilo escribe en su rango exclusivo.")
p("Se consideró usar un contador atómico compartido, que hubiera sido más corto de "
  "escribir, pero se descartó por dos razones. La primera es la contención, porque los "
  "12 hilos estarian peleando por la misma línea de cache. La segunda, y más "
  "importante, es que el orden de salida dependería de cual hilo llega primero y se "
  "perdería el determinismo. Con la suma de prefijos el resultado sale idéntico en "
  "cada corrida y además idéntico al de la versión secuencial, lo que permite "
  "comprobar que ambas versiones hacen el mismo trabajo.")
p("Los mecanismos de sincronía que se usan son cuatro. La cláusula reduction, que le "
  "da a cada hilo una copia privada y las combina al final. La barrera implícita al "
  "cerrar un parallel for, que garantiza que la rejilla de ocupación esté completa "
  "antes de que la siguiente etapa la lea. La directiva barrier explícita, que separa "
  "las tres fases del patrón de compactación. Y la directiva single, que hace que un "
  "solo hilo calcule la suma de prefijos mientras los demás esperan.")
p("Sobre el manejo de memoria, los arreglos de instancias se reservan una vez por "
  "mundo y se reutilizan en cada fotograma para no estar pidiendo memoria "
  "constantemente. Los objetos de OpenGL, que son el atlas de texturas, los "
  "renderizadores y el mapa de sombras, se crean al arrancar y se destruyen de forma "
  "ordenada al salir, incluso cuando el programa termina por un error.")

h1("Metodología de medición")
p("Para que la comparación entre las dos versiones tuviera sentido hizo falta "
  "asegurarse de que ambas hicieran exactamente el mismo trabajo. Los dos ejecutables "
  "comparten un modo de medición que se activa con el argumento --benchmark y que toma "
  "cinco decisiones, cada una para quitar de en medio una fuente de error distinta.")
vineta("Paso de tiempo fijo de un sesenta avo de segundo. Sin esto una máquina rápida "
       "avanzaría más la simulación en cada fotograma y terminaría haciendo más trabajo.")
vineta("Semilla fija. El mundo generado sale idéntico en cada repetición y entre las dos "
       "versiones.")
vineta("Sincronización vertical desactivada. Con ella activada el monitor limita todo a "
       "60 cuadros por segundo y la diferencia entre versiones queda escondida.")
vineta("Fotogramas de calentamiento descartados. Los primeros 120 fotogramas cargan la "
       "subida de texturas y el arranque del planificador de OpenMP.")
vineta("Llamada a glFinish antes de cerrar el cronómetro de cada fotograma. Sin esa "
       "espera el controlador gráfico encola trabajo y el tiempo medido saldría "
       "optimista.")
p("La prueba de que el control funcionó está en los informes. Las dos versiones "
  "reportan la misma cantidad de bloques generados y el mismo promedio de bloques "
  "dibujados. Si esos números no coincidieran, comparar los tiempos no tendría sentido.")
p("Las pruebas se corrieron en una computadora con procesador Intel Core i5-12400, que "
  "tiene 6 núcleos físicos y 12 hilos lógicos, con una ventana de 1280 por 720 píxeles "
  "y la semilla 20260901. Se probaron tres tamaños de N, que fueron 100 mil, 500 mil y "
  "2 millones, y cinco cantidades de hilos, que fueron 1, 2, 4, 8 y 12. Cada "
  "combinación se repitio 10 veces y cada repetición midió 300 fotogramas, para un "
  "total de 180 ejecuciones.")

# ============================================================================
#  RESULTADOS
# ============================================================================
h1("Resultados")
p("La Tabla 2 resume lo que se obtuvo con la mayor cantidad de hilos disponible. Los "
  "tiempos de CPU son el trabajo por fotograma que OpenMP reparte, sin contar el "
  "dibujo.")
tabla(
  ["N", "Bloques reales", "CPU secuencial", "CPU paralelo", "Speedup", "Eficiencia", "FPS sec.", "FPS par."],
  [
    ["100 000",   "95 874",    "0.51 ms",  "0.41 ms", "1.25x", "10.4 %", "1113", "1235"],
    ["500 000",   "482 184",   "2.99 ms",  "1.19 ms", "2.52x", "21.0 %", "232",  "393"],
    ["2 000 000", "1 918 719", "12.23 ms", "4.48 ms", "2.73x", "22.7 %", "53",   "117"],
  ],
  "Resultados con 12 hilos para los tres tamaños de carga probados.",
  anchos=[0.85, 1.0, 0.95, 0.85, 0.7, 0.8, 0.65, 0.65])
p("El mejor speedup del trabajo de CPU fue de 2.73 veces. Se alcanza con 8 hilos, "
  "donde la eficiencia todavía es de 34.2 por ciento, y ya no mejora al pasar a 12 "
  "hilos. Eso se explica porque la máquina tiene 6 núcleos físicos, y los hilos "
  "adicionales son hilos lógicos que comparten las unidades de ejecución del mismo "
  "núcleo, entonces no aportan un núcleo completo.")
figura("speedup_vs_hilos.png",
       "Speedup medido contra el speedup ideal lineal para los tres tamaños de N.", 5.6)
figura("eficiencia_vs_hilos.png",
       "Eficiencia paralela, que baja conforme se agregan hilos.", 5.6)
p("La eficiencia arranca en 73 por ciento con 2 hilos y baja hasta 23 por ciento con "
  "12. La caída es esperada, porque la parte que no se puede repartir pesa cada vez "
  "más en proporción y la sincronización crece. La lectura práctica es que el mejor "
  "aprovechamiento está entre 4 y 8 hilos, y que agregar más allá de 8 no mejora el "
  "tiempo y si desperdicia recursos.")
figura("tiempo_cpu_vs_hilos.png",
       "Tiempo de CPU por fotograma comparado con la línea base secuencial.", 6.3)
p("Hay un detalle en esas curvas que vale la pena resaltar. El punto de un solo hilo "
  "queda por encima de la línea base secuencial, o sea que la versión paralela "
  "corriendo con un hilo resulta entre 21 y 27 por ciento más lenta que la secuencial. "
  "Ese es el costo puro de OpenMP, porque el programa sigue pagando la creación del "
  "equipo de hilos, el reparto de iteraciones y las barreras, sin ningún hilo extra "
  "que lo compense.")

h2("Los tres speedups y por que son distintos")
p("El proyecto reporta tres speedups y cada uno mide una cosa distinta. Reportar solo "
  "el más alto daría una idea equivocada del trabajo.")
tabla(
  ["Que se mide", "Speedup con 12 hilos", "Explicación"],
  [
    ["Generación del mundo", "3.85x",
     "Es el mejor de los tres. Un bloque de trabajo grande y contiguo, sin sincronización con la tarjeta gráfica de por medio."],
    ["Trabajo de CPU por fotograma", "2.73x",
     "Corresponde a la parte que OpenMP reparte de verdad dentro del bucle principal."],
    ["Fotograma completo", "2.19x",
     "Es el más bajo porque incluye el dibujo, que es idéntico en las dos versiones y funciona como la fracción serial de la ley de Amdahl."],
  ],
  "Los tres speedups medidos con N igual a dos millones y su interpretación.",
  anchos=[1.7, 1.3, 3.4])
figura("generacion_vs_hilos.png",
       "Costo de generar un mundo completo, que es la etapa que mejor escala.", 6.3)
figura("etapas_por_fotograma.png",
       "Desglose del fotograma por etapa en los tres tamaños de carga.", 6.3)
p("En el desglose se ve que el armado del buffer domina sobre la física. Con dos "
  "millones de bloques la física cuesta 3.43 milisegundos y el descarte con "
  "empaquetado cuesta 8.80, o sea casi el triple. Saber eso antes de empezar evitó "
  "gastar esfuerzo optimizando la parte equivocada.")

h2("El caso que justifica el trabajo")
p("El resultado que mejor muestra para qué sirvió paralelizar aparece con dos millones "
  "de bloques. La versión secuencial promedia 53 cuadros por segundo y deja 264 de los "
  "300 fotogramas medidos por debajo de 60. La versión paralela sostiene 117 cuadros "
  "por segundo y prácticamente ningún fotograma queda por debajo de ese límite.")
figura("fps_vs_hilos.png",
       "Cuadros por segundo alcanzados frente al mínimo de 60 que pide el enunciado.", 6.3)
figura("fps_por_fotograma.png",
       "Ritmo fotograma a fotograma con dos millones de bloques, donde se ve que la mejora se sostiene.",
       5.8)
p("La segunda gráfica muestra que la ganancia no es un promedio afortunado. La línea "
  "de la versión secuencial se mantiene casi todo el tiempo debajo del límite y la de "
  "la versión paralela se mantiene estable alrededor de 120.")
figura("speedup_vs_n.png",
       "Como mejora el speedup conforme crece la carga de trabajo.", 5.6)
p("Al variar N se ve que el paralelismo rinde más cuando hay más trabajo que repartir. "
  "Con 100 mil bloques el speedup apenas llega a 1.25 veces, con 500 mil sube a 2.52 y "
  "con dos millones llega a 2.73. La razón es que el costo de abrir y cerrar las "
  "regiones paralelas es fijo, mientras que el trabajo útil crece con N, entonces ese "
  "costo pesa cada vez menos en proporción.")

# ============================================================================
#  CONCLUSIONES
# ============================================================================
salto()
h1("Conclusiones")
p("El trabajo cumplió el objetivo de tomar un programa secuencial y acelerarlo con "
  "OpenMP, y las mediciones respaldan cada afirmación.")
vineta("Se alcanzo un speedup de 2.73 veces en el trabajo de CPU por fotograma con 8 "
       "hilos y dos millones de bloques, con una eficiencia de 34.2 por ciento. La "
       "generación del mundo escala mejor y llega a 3.85 veces con 12 hilos.")
vineta("La paralelización fue lo que permitió sostener la experiencia de usuario que "
       "pide el enunciado. Con dos millones de elementos la versión secuencial deja 264 "
       "de 300 fotogramas debajo de 60 cuadros por segundo y la paralela prácticamente "
       "ninguno.")
vineta("La eficiencia baja conforme se agregan hilos, de 73 por ciento con 2 hilos a 23 "
       "por ciento con 12. El punto de mejor aprovechamiento quedó entre 4 y 8 hilos, y "
       "pasar de 8 no mejora nada porque la máquina tiene 6 núcleos físicos.")
vineta("Paralelizar no siempre conviene. Con 100 mil bloques la ganancia es marginal, y "
       "la versión paralela corriendo con un solo hilo resulta entre 21 y 27 por ciento "
       "más lenta que la secuencial por el costo fijo de OpenMP.")
vineta("Se logró eliminar las condiciones de carrera por diseño, sin usar critical, "
       "atomic ni locks en ninguna de las siete regiones paralelas. El patrón de "
       "compactación con suma de prefijos y el halo de la vegetación resolvieron los dos "
       "casos donde había riesgo real.")
vineta("Medir antes de optimizar sirvió. El desglose por etapa mostró que el descarte de "
       "bloques invisibles costaba casi el triple que la física, y eso guió hacia donde "
       "dirigir el esfuerzo.")

h1("Recomendaciones")
vineta("Paralelizar la actualización de los animales en caso de subir mucho su cantidad. "
       "Hoy son máximo 5000 entidades con lógica barata contra millones de bloques, "
       "entonces se dejo secuencial a propósito, pero con poblaciones mayores dejaría de "
       "ser despreciable.")
vineta("Probar un reparto dinámico en la vegetación. La cantidad de árboles por chunk "
       "varía según el relieve, entonces ahí si podría haber desbalance entre hilos, a "
       "diferencia del terreno donde el costo por columna es parejo.")
vineta("Repetir las mediciones en una máquina con más núcleos físicos para separar el "
       "efecto del hardware del efecto de la fracción serial. Con 6 núcleos no se puede "
       "distinguir del todo cual de los dos limita más.")
vineta("Reducir la fracción serial mandando más trabajo a la tarjeta gráfica, por "
       "ejemplo calculando el descarte de bloques invisibles con un shader de cómputo. "
       "Eso subiría el techo que hoy impone la ley de Amdahl.")

# ============================================================================
#  REFERENCIAS
# ============================================================================
salto()
h1("Referencias bibliográficas")
def ref(texto):
    par = p(texto, 11)
    par.paragraph_format.left_indent = Inches(0.5)
    par.paragraph_format.first_line_indent = Inches(-0.5)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_after = Pt(10)

ref("Amdahl, G. M. (1967). Validity of the single processor approach to achieving "
    "large scale computing capabilities. AFIPS Conference Proceedings, 30, 483-485. "
    "Reimpreso en IEEE Solid-State Circuits Society Newsletter, 12(3), 19-20, 2007.")
ref("Chapman, B., Jost, G. y van der Pas, R. (2007). Using OpenMP. Portable Shared "
    "Memory Parallel Programming. Cambridge, Massachusetts. The MIT Press.")
ref("Foster, I. (1995). Designing and Building Parallel Programs. Concepts and Tools "
    "for Parallel Software Engineering. Reading, Massachusetts. Addison-Wesley.")
ref("OpenMP Architecture Review Board. (2021). OpenMP Application Programming "
    "Interface. Versión 5.2. Recuperado de https://www.openmp.org/specifications/")

# ============================================================================
#  Apéndice
# ============================================================================
salto()
h1("Apéndice. Material suplementario")
p("El repositorio del código fuente incluye material que sirve para reproducir todo lo "
  "que aparece en este informe.")
tabla(
  ["Ruta en el repositorio", "Que contiene"],
  [
    ["Secuencial/secuencial.cpp", "Versión secuencial completa, sin ninguna directiva de OpenMP."],
    ["Paralelo/paralelo.cpp", "Versión paralela con las siete regiones de OpenMP."],
    ["common/", "Código compartido por las dos versiones, como el atlas de texturas, los renderizadores y la captura de argumentos."],
    ["common/benchmark.cpp", "Modo de medición compartido que produce los informes."],
    ["metricas/medir_y_graficar.py", "Script que compila, corre las pruebas, arma los CSV y dibuja las gráficas."],
    ["metricas/resultados/mediciones_crudas.csv", "Las 180 ejecuciones con todas sus repeticiones."],
    ["metricas/resultados/resumen.csv", "Promedios, speedup y eficiencia por combinación de N e hilos."],
    ["metricas/resultados/", "Las gráficas que aparecen en este informe."],
    ["docs/pipeline/", "Imagenes del pipeline del programa."],
    ["README.md", "Instrucciones para compilar, correr y medir, con la lista completa de argumentos."],
  ],
  "Material suplementario disponible en el repositorio del código.",
  anchos=[2.4, 4.0])
p("Para reproducir las mediciones basta compilar el proyecto y correr el script de "
  "métricas, que se encarga del resto.")
par = p("python metricas/medir_y_graficar.py", 10)
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
par.runs[0].font.name = "Consolas"
p("El script acepta argumentos para cambiar los valores de N, la cantidad de hilos, "
  "las repeticiones y los fotogramas medidos.")

# ============================================================================
#  ANEXO 1
# ============================================================================
salto()
h1("Anexo 1. Diagrama de flujo del programa")
p("El diagrama recorre el programa completo, desde que recibe los argumentos hasta que "
  "libera la memoria. Los colores separan lo que corre en un solo hilo, lo que se "
  "reparte con OpenMP, los mecanismos de sincronía y lo que ejecuta la tarjeta "
  "gráfica.")
figura("anexo1_diagrama_flujo.png",
       "Diagrama de flujo del programa con las secciones paralelas y los mecanismos de sincronía señalados.",
       6.3)
p("Los seis puntos que pide el enunciado aparecen marcados en el diagrama. La captura "
  "de argumentos está al inicio, en la primera columna. La solicitud de ingreso de "
  "datos ocurre por línea de comandos, donde N es obligatorio y el resto de opciones "
  "traen valor por omisión. La programación defensiva es la rama que se toma cuando un "
  "argumento sale inválido o fuera de rango, y también la que atiende los errores al "
  "buscar las texturas o al crear la ventana. Las secciones paralelas son las cajas "
  "anaranjadas. Los mecanismos de sincronía son las cajas moradas, que incluyen la "
  "barrera implícita del parallel for y las directivas barrier y single. El despliegue "
  "de resultados aparece dos veces, en el título de la ventana durante la ejecución y "
  "en el informe final cuando se corre en modo de medición.")

# ============================================================================
#  ANEXO 2
# ============================================================================
salto()
h1("Anexo 2. Catálogo de funciones")
p("Se documentan las funciones donde ocurre el trabajo que crece con N y las que "
  "sostienen la infraestructura del programa. Los tipos que aparecen son los del "
  "código. World es la estructura del mundo, AppConfig es la configuración leída de la "
  "línea de comandos, BlockId es el tipo de bloque e InstanceData es el registro de 20 "
  "bytes que se envía a la tarjeta gráfica.")

h2("Generación del mundo")
tabla(
  ["Función", "Entradas", "Salidas", "Descripción"],
  [
    ["parseArguments",
     "argc de tipo int, argv de tipo char**, programName de tipo const char*",
     "config de tipo AppConfig, error de tipo std::string, retorna bool",
     "Lee los argumentos de la línea de comandos. Reconoce N como valor posicional y el resto como opciones con nombre. Rechaza valores no numericos, fuera de rango, opciones desconocidas y opciones sin su valor. Devuelve false y llena error cuando algo está mal."],
    ["surfaceHeight",
     "gx y gz de tipo int, seed de tipo uint32_t, relief de tipo float, biome de tipo const Biome&",
     "Retorna int con el índice Y del bloque más alto de esa columna",
     "Suma varias octavas de ruido para obtener colinas con detalle a distintas escalas. Es una función pura, no depende de estado global, entonces se puede evaluar en paralelo sobre columnas distintas."],
    ["generateVoxels",
     "world de tipo World&, relief de tipo float",
     "Retorna std::vector<BlockId> con la retícula, y llena world.heightMap y world.height",
     "Recorre las columnas del plano y decide el tipo de cada celda. Es el primer punto caliente. Se divide en tres pasos, que son el mapa de alturas, los estratos y la vegetación."],
    ["plantVegetation",
     "voxels de tipo std::vector<BlockId>&, world de tipo const World&",
     "Modifica voxels colocando troncos, hojas y cactus",
     "Coloca los árboles sobre las columnas elegidas. Divide el plano en chunks 2D con un chunk por hilo y usa un halo de dos celdas. Cada hilo lee el halo pero escribe únicamente dentro de su chunk."],
    ["buildBlockList",
     "voxels de tipo const std::vector<BlockId>&, world de tipo World&, cfg de tipo const AppConfig&",
     "Llena world.blocks y world.occupancy",
     "Convierte la retícula en la lista compacta de bloques que se anima y se dibuja. Asigna a cada bloque el instante en que aparece y en que se desvanece. Usa suma de prefijos para conservar el mismo orden que la versión secuencial."],
    ["generateWorld",
     "cfg de tipo const AppConfig&, seed de tipo uint32_t",
     "world de tipo World& queda generado, retorna double con los milisegundos que tomó",
     "Orquesta la generación completa. Elige el bioma, deduce el lado de la retícula que aproxima los N bloques pedidos y llama a las funciones anteriores."],
  ],
  "Funciones que participan en la generación del mundo.",
  anchos=[1.0, 1.6, 1.6, 2.2])

h2("Actualización y dibujo por fotograma")
tabla(
  ["Función", "Entradas", "Salidas", "Descripción"],
  [
    ["updateWorld",
     "world de tipo World&, elapsed de tipo float, dt de tipo float",
     "liveBlocks de tipo size_t& con los bloques que siguen vivos",
     "Avanza la máquina de estados y la física de cada bloque. Es el segundo punto caliente. Los bloques que aterrizan marcan su celda como ocupada y los que se desvanecen la liberan. La barrera implícita al cerrar la región ordena esta etapa antes del descarte."],
    ["blockIsVisible",
     "world de tipo const World&, b de tipo const Block&",
     "Retorna bool",
     "Decide si un bloque puede verse. Devuelve false cuando el bloque ya desapareció o cuando sus seis vecinos están ocupados, porque en ese caso ninguna de sus caras es observable."],
    ["buildInstanceBuffer",
     "world de tipo const World&, faceTex de tipo const uint32_t*, instances de tipo InstanceData*, visibleFlags de tipo uint8_t*",
     "Retorna size_t con la cantidad de instancias escritas",
     "Descarta los bloques invisibles y empaqueta el resto en el formato que consume la tarjeta gráfica. Es el tercer punto caliente y el más caro. Usa suma de prefijos para que cada hilo escriba en un rango exclusivo."],
    ["updateMobsSequential",
     "mobs de tipo std::vector<Mob>&, world de tipo const World&, dt de tipo float",
     "Modifica el estado de cada animal",
     "Avanza el ciclo de vida y la inteligencia artificial de los animales. Cada uno decide entre esperar o caminar, y al caminar avanza usando coseno y seno de su ángulo. Verifica que la celda destino no tenga árbol ni cactus y que el desnivel no pase de un bloque."],
    ["buildNearbyShadowInstances",
     "source de tipo const InstanceData*, count de tipo size_t, focus de tipo const glm::vec3&, radius de tipo float",
     "output de tipo std::vector<InstanceData>&, retorna size_t con la cantidad",
     "Selecciona los bloques cercanos a la cámara que pueden proyectar sombra, para no volver a enviar el terreno lejano completo. Usa el mismo patrón de compactación con suma de prefijos."],
    ["calculateCamera",
     "world de tipo const World&, cfg de tipo const AppConfig&, now de tipo double",
     "Retorna CameraView con la posición, el objetivo y el campo de visión",
     "Calcula el encuadre activo de la cámara. En modo órbita mantiene la toma exterior original y en modo automático alterna entre cinco encuadres con transiciones suaves. Usa seno y coseno para los recorridos."],
  ],
  "Funciones que se ejecutan en cada fotograma.",
  anchos=[1.1, 1.6, 1.5, 2.2])

h2("Infraestructura y medición")
tabla(
  ["Función o clase", "Entradas", "Salidas", "Descripción"],
  [
    ["CubeRenderer::draw",
     "instances de tipo const InstanceData*, count de tipo size_t, viewProj de tipo const glm::mat4&, atlasTex de tipo GLuint, lighting de tipo const SceneLighting&",
     "Dibuja en el buffer activo de OpenGL",
     "Transfiere el arreglo de instancias y emite una sola llamada de dibujo instanciado para todo el mundo. Debe llamarse siempre desde el hilo que posee el contexto de OpenGL."],
    ["TextureAtlas::build",
     "assetsDir de tipo const std::string&",
     "error de tipo std::string&, retorna bool",
     "Decodifica los archivos PNG, aplica los tintes por bioma y arma el bloque de píxeles que se sube a la tarjeta gráfica. No toca OpenGL, entonces se puede probar aparte."],
    ["Benchmark::start",
     "cfg de tipo const AppConfig&",
     "Deja el cronómetro configurado",
     "Copia los parámetros de medición. Si la cantidad de fotogramas a medir es cero el objeto queda inactivo y el programa se comporta como protector de pantalla normal."],
    ["Benchmark::addFrame",
     "timing de tipo const FrameTiming&, drawnBlocks de tipo size_t",
     "Acumula la muestra",
     "Descarta en silencio los fotogramas de calentamiento y guarda el resto hasta completar la cantidad pedida."],
    ["Benchmark::report",
     "info de tipo const BenchmarkInfo&",
     "error de tipo std::string&, retorna bool, imprime el informe",
     "Calcula promedios, minimos, maximos y percentiles, e imprime el informe en formato clave igual valor. Si se pidio el CSV por fotograma también lo escribe."],
    ["elapsedMsAndReset",
     "mark de tipo std::chrono::steady_clock::time_point&",
     "Retorna double con los milisegundos y adelanta mark",
     "Cronometra una etapa del fotograma y deja la marca lista para la siguiente, para no repetir la misma expresión en cada etapa."],
  ],
  "Funciones de infraestructura y del modo de medición.",
  anchos=[1.2, 1.7, 1.4, 2.1])

# ============================================================================
#  ANEXO 3
# ============================================================================
salto()
h1("Anexo 3. Bitácora de pruebas")
p("Se corrieron 180 ejecuciones en total. Cada combinación de N y de cantidad de hilos "
  "se repitio 10 veces, y cada repetición midió 300 fotogramas después de descartar "
  "120 de calentamiento. Las tablas de esta sección vienen del archivo "
  "mediciones_crudas.csv que genera el script de métricas.")

import csv as _csv
CRUDAS = Path(r"d:/Tareas/Computacion Paralela/Proyecto 1/metricas/resultados/mediciones_crudas.csv")
with CRUDAS.open(encoding="utf-8") as f:
    filas_csv = list(_csv.DictReader(f))

def repeticiones(ver, hilos, n):
    return [r for r in filas_csv
            if r["version"] == ver and r["hilos"] == str(hilos)
            and r["n_solicitado"] == str(n)]

def tabla_repeticiones(ver, hilos, n, titulo):
    reps = repeticiones(ver, hilos, n)
    filas = []
    for r in reps:
        filas.append([r["repeticion"], r["update_ms_prom"], r["build_ms_prom"],
                      r["cpu_ms_prom"], r["frame_ms_prom"], r["fps_prom"],
                      r["frames_bajo_60fps"]])
    prom = lambda c: sum(float(r[c]) for r in reps) / len(reps)
    filas.append(["Promedio", "%.4f" % prom("update_ms_prom"), "%.4f" % prom("build_ms_prom"),
                  "%.4f" % prom("cpu_ms_prom"), "%.4f" % prom("frame_ms_prom"),
                  "%.2f" % prom("fps_prom"), "%.1f" % prom("frames_bajo_60fps")])
    tabla(["Rep.", "update ms", "build ms", "CPU ms", "Fotograma ms", "FPS", "Bajo 60 FPS"],
          filas, titulo, anchos=[0.7, 0.95, 0.95, 0.95, 1.15, 0.95, 0.95])
    return prom("cpu_ms_prom")

h2("Prueba con dos millones de bloques")
p("Esta es la prueba principal, porque es la única carga donde la versión secuencial "
  "no logra sostener los 60 cuadros por segundo. El mundo generado tiene 1 918 719 "
  "bloques.")
cpu_sec_2m = tabla_repeticiones("secuencial", 1, 2000000,
  "Diez repeticiones de la versión secuencial con N igual a dos millones.")
cpu_par_2m = tabla_repeticiones("paralelo", 8, 2000000,
  "Diez repeticiones de la versión paralela con 8 hilos y N igual a dos millones.")
p("El promedio del trabajo de CPU por fotograma pasó de %.3f milisegundos en la "
  "versión secuencial a %.3f en la paralela con 8 hilos. Eso da un speedup de %.2f "
  "veces y una eficiencia de %.1f por ciento. En la columna de la derecha se ve la "
  "diferencia más importante, porque la versión secuencial deja alrededor de 264 "
  "fotogramas de 300 por debajo de 60 cuadros por segundo y la paralela no deja "
  "ninguno."
  % (cpu_sec_2m, cpu_par_2m, cpu_sec_2m / cpu_par_2m,
     cpu_sec_2m / cpu_par_2m / 8 * 100))

h2("Prueba con quinientos mil bloques")
cpu_sec_500 = tabla_repeticiones("secuencial", 1, 500000,
  "Diez repeticiones de la versión secuencial con N igual a quinientos mil.")
cpu_par_500 = tabla_repeticiones("paralelo", 8, 500000,
  "Diez repeticiones de la versión paralela con 8 hilos y N igual a quinientos mil.")
p("Con esta carga el speedup fue de %.2f veces y la eficiencia de %.1f por ciento con "
  "8 hilos. Las dos versiones se mantienen arriba de 60 cuadros por segundo, entonces "
  "aquí la mejora se nota en el margen disponible y no en la experiencia de uso."
  % (cpu_sec_500 / cpu_par_500, cpu_sec_500 / cpu_par_500 / 8 * 100))

h2("Prueba con cien mil bloques")
cpu_sec_100 = tabla_repeticiones("secuencial", 1, 100000,
  "Diez repeticiones de la versión secuencial con N igual a cien mil.")
cpu_par_100 = tabla_repeticiones("paralelo", 8, 100000,
  "Diez repeticiones de la versión paralela con 8 hilos y N igual a cien mil.")
p("Aquí el speedup baja a %.2f veces y la eficiencia a %.1f por ciento. El trabajo por "
  "fotograma es tan poco que el costo fijo de abrir y cerrar las regiones paralelas se "
  "come casi toda la ganancia."
  % (cpu_sec_100 / cpu_par_100, cpu_sec_100 / cpu_par_100 / 8 * 100))

h2("Resumen de speedup y eficiencia")
p("La Tabla 12 junta todas las combinaciones probadas. El speedup se calcula dividiendo "
  "el tiempo promedio de la versión secuencial entre el de la paralela, y la eficiencia "
  "es ese speedup dividido entre la cantidad de hilos.")

RESUMEN = Path(r"d:/Tareas/Computacion Paralela/Proyecto 1/metricas/resultados/resumen.csv")
with RESUMEN.open(encoding="utf-8") as f:
    res = list(_csv.DictReader(f))
filas_res = []
for r in res:
    filas_res.append([
        "{:,}".format(int(r["n_solicitado"])).replace(",", " "),
        r["hilos"],
        "%.3f" % float(r["cpu_secuencial_ms"]),
        "%.3f" % float(r["cpu_paralelo_ms"]),
        "%.2fx" % float(r["speedup_cpu"]),
        "%.1f %%" % float(r["eficiencia_pct"]),
        "%.1f" % float(r["fps_secuencial"]),
        "%.1f" % float(r["fps_paralelo"]),
    ])
tabla(["N", "Hilos", "CPU sec. ms", "CPU par. ms", "Speedup", "Eficiencia", "FPS sec.", "FPS par."],
      filas_res,
      "Speedup y eficiencia de las quince combinaciones de N y cantidad de hilos.",
      anchos=[0.85, 0.6, 0.95, 0.95, 0.8, 0.9, 0.75, 0.75])

h2("Captura de las mediciones")
p("La siguiente captura muestra el informe que imprime cada ejecución al terminar de "
  "medir, y la tabla resumen que arma el script después de recorrer las 180 corridas.")
figura("anexo3_captura_mediciones.png",
       "Salida del programa en modo de medición y resumen del script de métricas.", 6.3)
p("El informe por ejecución reporta los tiempos de cada etapa, los cuadros por segundo "
  "promedio y mínimos, y cuántos fotogramas quedaron por debajo de 60. Ese último dato "
  "es el que conecta la medición con lo que pide el enunciado.")

# ============================================================================
#  GUARDAR
# ============================================================================
salida = Raíz / "Informe - Proyecto 1 Computacion Paralela.docx"
doc.save(str(salida))
print("Documento guardado en", salida)
print("Tablas:", CONT["tabla"], " Figuras:", CONT["figura"])
